from collections import ChainMap
import time

from docx import Document
from docx.shared import Inches

from habit import Habit, WeekPeriod, DaysOfWeek

fruit_prices = {
    'apples': 98,
    'oranges': 110,
    'bananas': 120,
    'pineapples': 60,
}

vegetable_prices = {
    'tomatoes': 70,
    'cucumbers': 40,
    'pineapples': 888,
}

assortment = ChainMap(fruit_prices, vegetable_prices)

print(assortment)
print(assortment.values())
print(assortment.keys())

for key, item in assortment.items():
    print(key, item)  ## passes repeated key 'pineapples'


############################################################

def function_to_be_tested(param: int):
    return f'You have entered number: {param}'


############################################################

def test_create_document():
    document = Document()

    document.add_heading('Document Title', 0)

    p = document.add_paragraph('A plain paragraph having some ')
    p.add_run('bold').bold = True
    p.add_run(' and some ')
    p.add_run('italic.').italic = True

    document.add_heading('Heading, level 1', level=1)
    document.add_paragraph('Intense quote', style='Intense Quote')

    document.add_paragraph(
        'first item in unordered list', style='List Bullet'
    )
    document.add_paragraph(
        'first item in ordered list', style='List Number'
    )

    document.add_picture('./q101.jpg', width=Inches(1.25))

    records = (
        (3, '101', 'Spam'),
        (7, '422', 'Eggs'),
        (4, '631', 'Spam, spam, eggs, and spam')
    )

    table = document.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Qty'
    hdr_cells[1].text = 'Id'
    hdr_cells[2].text = 'Desc'
    for qty, id, desc in records:
        row_cells = table.add_row().cells
        row_cells[0].text = str(qty)
        row_cells[1].text = id
        row_cells[2].text = desc

    document.add_page_break()

    document.save('demo.docx')


############################################################

my_habit = Habit(name='Reading',
                 description='Habit of reading 10 pages of professional literature every day')
my_habit.set_precondition('After work')
my_habit.set_place('At work table')

my_habit.set_period(
    WeekPeriod(days_of_week=[DaysOfWeek.MONDAY,
                             DaysOfWeek.TUESDAY,
                             DaysOfWeek.WEDNESDAY,
                             DaysOfWeek.THURSDAY,
                             DaysOfWeek.FRIDAY],
               time = '21:00')
)
my_habit.set_period(
    WeekPeriod(days_of_week=[DaysOfWeek.SATURDAY,
                             DaysOfWeek.SUNDAY],
               time = '11:00')
)

print(my_habit)
