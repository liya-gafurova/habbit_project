import datetime
import calendar
import math
from collections import namedtuple
from enum import Enum

from docx import Document
from docx.shared import Inches, RGBColor
from docx.enum.section import WD_ORIENTATION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.styles.styles import Styles

from app.fillers import TABLE_FILLERS
from app.habit import Habit, Months, WeekPeriod, DaysOfWeek

DocumentDaysOfWeek = dict(zip([1,2,3,4,5,6,7,],
                              ['M', 'T', 'W', 'T', 'F', 'S', 'S']))



def test_create_document():
    document = Document()

    document.add_heading('Document Title', 0)

    p = document.add_paragraph('A plain paragraph having some ')
    p.add_run('bold').bold = True
    p.add_run(' and some ')
    p.add_run('italic.').italic = True

    document.add_heading('Heading, level 1', level=1)
    document.add_paragraph('Intense quote', style='Intense Quote')

    document.add_paragraph(
        'first item in unordered list', style='List Bullet'
    )
    document.add_paragraph(
        'first item in ordered list', style='List Number'
    )

    document.add_picture('./q101.jpg', width=Inches(1.25))

    records = (
        (3, '101', 'Spam'),
        (7, '422', 'Eggs'),
        (4, '631', 'Spam, spam, eggs, and spam')
    )

    table = document.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Qty'
    hdr_cells[1].text = 'Id'
    hdr_cells[2].text = 'Desc'
    for qty, id, desc in records:
        row_cells = table.add_row().cells
        row_cells[0].text = str(qty)
        row_cells[1].text = id
        row_cells[2].text = desc

    document.add_page_break()

    document.save('demo.docx')


class HabitDocument:
    CURRENT_YEAR = datetime.datetime.today().year
    NUMBER_OF_DAYS_IN_ROW = 7  # week in row

    def __init__(self, name, description, preconditions, place, week_periods, month, year=None):
        self.name_of_habit = name
        self.habit_description = description
        self.preconditions = preconditions
        self.place = place
        self.week_periods:WeekPeriod = week_periods
        self.month = month
        self.year = year if year else self.CURRENT_YEAR

        self.document: Document = Document()

    @property
    def days_of_week_over_month_distribution(self):

        days_of_week_over_month_distribution = {}
        for day in range(1, self.days_in_month + 1):
            week_day_number = datetime.datetime.isoweekday(datetime.datetime(self.year, self.month, day))
            time = ''
            for week_period in self.week_periods:
                if week_day_number in week_period.days_of_week:
                    time = week_period.time
            days_of_week_over_month_distribution[day]= (week_day_number, time)

        return days_of_week_over_month_distribution



    def create_document(self, document_name):
        self._create_document_heading()
        self._create_table()
        self.document.save(document_name)

    def _create_document_heading(self):
        # landscape orientation
        section = self.document.sections[0]
        new_width, new_height = section.page_height, section.page_width
        section.orientation = WD_ORIENTATION.LANDSCAPE
        section.page_width = new_width
        section.page_height = new_height

        # Heading - name of Habit
        heading = self.document.add_heading(self.name_of_habit)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        line = self.document.add_paragraph('Description: ')
        line.add_run(self.habit_description).bold = True

        line = self.document.add_paragraph('Preconditions: ')
        for precondition in self.preconditions:
            line.add_run(precondition).bold = True
            line.add_run(', ').bold = True

        line = self.document.add_paragraph('Place: ')
        line.add_run(self.place).bold = True

    def _create_table(self):
        rows, columns = self._calculate_table_size()
        table = self.document.add_table(rows=rows, cols=columns)
        day_counter = 1
        zero_month_filler = 0 if self.month < 10 else ''


        for row in table.rows:
            for cell in row.cells:
                space_fillers = ' '*9 if day_counter > 9 else ' '*11
                day_of_week = DocumentDaysOfWeek[self.days_of_week_over_month_distribution[day_counter][0]]
                time = self.days_of_week_over_month_distribution[day_counter][1]

                cell_header = cell.paragraphs[0].add_run(f"{day_counter}/{zero_month_filler}{self.month}{space_fillers}{day_of_week}")
                cell_header.font.color.rgb = RGBColor(120, 120, 120)

                habit_time_line = cell.add_paragraph().add_run(" "*6 + f'{time}\n')
                habit_time_line.font.bold = True





                if day_counter == self.days_in_month:
                    break
                day_counter += 1

    def _calculate_table_size(self):
        _, self.days_in_month = calendar.monthrange(self.year, self.month)
        rows = math.ceil(self.days_in_month / self.NUMBER_OF_DAYS_IN_ROW)
        columns = self.NUMBER_OF_DAYS_IN_ROW

        return rows, columns



my_habit = Habit(name='Reading',
                 description='Habit of reading 10 pages of professional literature every day')
my_habit.set_precondition('After work')
my_habit.set_place('At work table')
my_habit.set_period(WeekPeriod(
    days_of_week=[DaysOfWeek.SATURDAY, DaysOfWeek.SUNDAY],
    time='8:00'
))
my_habit.set_period(WeekPeriod(
    days_of_week=[DaysOfWeek.MONDAY, DaysOfWeek.THURSDAY],
    time='8:15'
))

# TODO just Habit Object
doc = HabitDocument(
    name=my_habit.name,
    description=my_habit.description,
    preconditions=my_habit.preconditions,
    place=my_habit.place,

    week_periods=my_habit.periods,
    month=Months.AUGUST
)
doc.create_document(f'../files/first_habit_doc_{datetime.datetime.now()}.docx')
