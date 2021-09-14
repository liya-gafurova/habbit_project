import math

from docx import Document
from docx.shared import Inches, RGBColor
from docx.enum.section import WD_ORIENTATION
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.domain.habitentity import HabitEntity
from app.presenters.data_extractor import DataExtractor, OneDay


class HabitDocument:
    NUMBER_OF_DAYS_IN_ROW = 7  # week in row

    def __init__(self, habit: HabitEntity):
        self.data_extractor = DataExtractor(habit)
        self.document: Document = Document()

    def create_document(self, document_name):
        self.make_landscape_orientation()
        self.make_description_info()
        self.make_table()
        self.document.save(document_name)

    def make_landscape_orientation(self):
        # landscape orientation
        section = self.document.sections[0]
        new_width, new_height = section.page_height, section.page_width
        section.orientation = WD_ORIENTATION.LANDSCAPE
        section.page_width = new_width
        section.page_height = new_height

    def make_description_info(self):
        # TODO capitalize
        header_data = self.data_extractor.get_header_content()

        heading = self.document.add_heading(header_data.name)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        line = self.document.add_paragraph('Description: ')
        line.add_run(header_data.description).bold = True

        line = self.document.add_paragraph('Preconditions: ')
        line.add_run(header_data.preconditions).bold = True

        line = self.document.add_paragraph('Place: ')
        line.add_run(header_data.place).bold = True

    def make_table(self):

        rows, columns = self._calculate_table_size(self.data_extractor.days_in_month)
        table = self.document.add_table(rows=rows, cols=columns)

        self.fill_table(table)

    def fill_table(self, table):
        day_counter = 1
        for row in table.rows:
            for cell in row.cells:
                one_day = self.data_extractor.planned_month[day_counter]
                self.add_cell_header(cell, one_day)
                self.add_cell_time_paragraph(cell, one_day)

                if day_counter == self.data_extractor.days_in_month:
                    break
                day_counter += 1

    def add_cell_header(self, cell, day: OneDay):
        space_fillers = self._calculate_number_of_spaces(day.day_number)
        zero_month_filler = 0 if day.month < 10 else ''  # data extractors
        cell_header = cell.paragraphs[0].add_run(
            f"{day.day_number}/{zero_month_filler}{day.month}{space_fillers}{day.weekday_letter}")
        cell_header.font.color.rgb = RGBColor(120, 120, 120)

    def add_cell_time_paragraph(self, cell, day: OneDay):
        habit_time_line = cell.add_paragraph().add_run(" " * 6 + f'{day.time}\n')
        habit_time_line.font.bold = True

    def _calculate_table_size(self, days_in_month):
        rows = math.ceil(days_in_month / self.NUMBER_OF_DAYS_IN_ROW)
        columns = self.NUMBER_OF_DAYS_IN_ROW
        return rows, columns

    def _calculate_number_of_spaces(self, day_of_month):
        return ' ' * 9 if day_of_month > 9 else ' ' * 11
