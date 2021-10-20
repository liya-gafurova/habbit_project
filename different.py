import copy
import json
import logging
import random
import re
import typing
from dataclasses import dataclass, asdict
from functools import reduce
from typing import NamedTuple
from collections import namedtuple, ChainMap, Counter

import pafy
import pytest
from docx import Document

logger = logging.getLogger(__name__)


@dataclass(frozen=True)
class Name:
    first_name: str
    surname: str


class Money(NamedTuple):
    currency: str
    value: int


Line = namedtuple('Line', ['sku', 'qty'])


def test_equality():
    # objects with equal fields are equal
    assert Money('gbp', 10) == Money('gbp', 10)
    assert Name('Harry', 'Percival') != Name('Bob', 'Gregory')
    assert Line('RED-CHAIR', 5) == Line('RED-CHAIR', 5)


test_equality()

fiver = Money('gbp', 5)
tenner = Money('gbp', 10)


def can_add_money_values_for_the_same_currency():
    assert fiver + fiver == tenner


def can_subtract_money_values():
    assert tenner - fiver == fiver


def adding_different_currencies_fails():
    with pytest.raises(ValueError):
        Money('usd', 10) + Money('gbp', 10)


def can_multiply_money_by_a_number():
    assert fiver * 5 == Money('gbp', 25)


def multiplying_two_money_values_is_an_error():
    with pytest.raises(TypeError):
        tenner * fiver


fruit_prices = {
    'apples': 98,
    'oranges': 110,
    'bananas': 120,
    'pineapples': 60,
}

vegetable_prices = {
    'tomatoes': 70,
    'cucumbers': 40,
    'pineapples': 888,
}

assortment = ChainMap(fruit_prices, vegetable_prices)

print(assortment)
print(assortment.values())
print(assortment.keys())

for key, item in assortment.items():
    print(key, item)  ## passes repeated key 'pineapples'


############################################################

def function_to_be_tested(param: int):
    return f'You have entered number: {param}'


############################################################

def test_create_document():
    document = Document()

    document.add_heading('Document Title', 0)

    p = document.add_paragraph('A plain paragraph having some ')
    p.add_run('bold').bold = True
    p.add_run(' and some ')
    p.add_run('italic.').italic = True

    document.add_heading('Heading, level 1', level=1)
    document.add_paragraph('Intense quote', style='Intense Quote')

    document.add_paragraph(
        'first item in unordered list', style='List Bullet'
    )
    document.add_paragraph(
        'first item in ordered list', style='List Number'
    )

    document.add_picture('./q101.jpg', width=Inches(1.25))

    records = (
        (3, '101', 'Spam'),
        (7, '422', 'Eggs'),
        (4, '631', 'Spam, spam, eggs, and spam')
    )

    table = document.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Qty'
    hdr_cells[1].text = 'Id'
    hdr_cells[2].text = 'Desc'
    for qty, id, desc in records:
        row_cells = table.add_row().cells
        row_cells[0].text = str(qty)
        row_cells[1].text = id
        row_cells[2].text = desc

    document.add_page_break()

    document.save('demo.docx')


############################################################

## abstract protocol that describes behavior of an paticular class

# class TransformationParamsToDocumentUseCase(typing.Protocol):
#     def transform(self, params, document):
#         pass
#
#
# class TransformHabitToDocument(TransformationParamsToDocumentUseCase):
#     def transform(self, params: habit.HabitEntity, document: habit.DocxDocument):
#         pass


## Polymorphism

class Cat:
    def __init__(self, name):
        self.name = name
        self.sound = 'Purrrr'

    def make_sound(self):
        print(self.sound)


class Dog:
    def __init__(self, name):
        self.name = name
        self.sound = 'Auff'

    def make_sound(self):
        print(self.sound)


animals = [Dog('Bobic'), Cat('Marysa'), Dog('Jack')]
for animal in animals:
    animal.make_sound()

# try lambdas

# сделать какое-то действие над каждым элементом списка
l = [46, 4, 11, 12, 64, 59, 59, 96, 39, 61, 80, 12, 32, 30, 14, 38, 53, 72, 11, 16]
l2 = list(map(lambda x: x // 10, l))

# https://tproger.ru/problems/python-3-exercises-for-beginners-geekbrains/
# Задача 22
# Напишите программу, которая принимает текст и выводит два слова: наиболее часто встречающееся и самое длинное.
text = '''In this situation, both the normal function and the lambda behave similarly. In the next section, you’ll see a situation where the behavior of a lambda can be deceptive due to its evaluation time (definition time vs runtime).'''

words = re.sub(re.compile('[.,!?-`\'’()]'), '', text).split(' ')
word_lens = list(map(len, words))
word_counts = Counter(words)
longest = words[word_lens.index(max(word_lens))]
frequent = next(
    iter({k: v for k, v in sorted(word_counts.items(), key=lambda item: item[1], reverse=True)}))  # sort dict by value
print(longest, frequent)

## ----
word_counts = Counter(words)
frequent = word_counts.most_common()[0]
longest = max(words, key=len)
print(longest, frequent)

# С помощью анонимной функции (!!! lambda) извлеките (!!! filter()) из списка числа (list), делимые на 15.
l = [46, 4, 11, 12, 64, 59, 59, 96, 39, 30, 61, 80, 12, 32, 30, 14, 38, 53, 72, 11, 16]
ll = list(filter(lambda element: element % 15 == 0, l))
# https://acmp.ru/index.asp?main=tasks&str=%20&page=19&id_type=0 -- олимпиадные  задачи по программированию
# map / reduce

# map ()
# filter()
# reduce()

numbers = [3, 4, 5, 6]


def cum_sum(first, second):
    return first + second


d = reduce(cum_sum, numbers)
print(d)


# decorators

def wrapper(func):
    def inner(list_range):
        print('before function')
        func(list_range)
        print('after_function')

    return inner


@wrapper
def test_func(list_range):
    print([random.randint(1, list_range) for i in range(list_range)])


test_func(5)

## interpolation

from scipy.interpolate import interp1d
import numpy as np
import matplotlib.pyplot as plt

# TODO how to make commit in the past

time_start = 123
time_max_value = 158
max_value = 10000
n = 5


def generate_intermediate_points(start_point, end_point, N, intermediate_points=[]):
    if len(intermediate_points) >= N:
        return intermediate_points
    else:
        possible_points = np.linspace(start_point, end_point, 10, endpoint=False)
        point = random.choice(possible_points[1:])
        intermediate_points.append(point)
        return generate_intermediate_points(intermediate_points[-1], end_point, N, intermediate_points)


x = np.array([time_start] + generate_intermediate_points(time_start, time_max_value, n, []) + [time_max_value])
y = np.array([0] + generate_intermediate_points(0, max_value, n, []) + [max_value])

# he string has to be one of ‘linear’, ‘nearest’, ‘nearest-up’,
# ‘zero’, ‘slinear’, ‘quadratic’, ‘cubic’, ‘previous’, or ‘next’.
f1 = interp1d(x, y, kind='linear')
plt.plot(x, y, 'o', x, f1(x), '-')
plt.show()


def get_accumulation_function2(fight_duration_seconds, max_value_seconds, max_value):
    x = [0, max_value_seconds, fight_duration_seconds]
    y = [0, max_value, max_value]
    linear_spline_func = interp1d(x, y, kind='linear')

    plt.plot(x, y, 'o', x, linear_spline_func(x), '-')
    plt.plot([max_value_seconds ], [linear_spline_func(max_value_seconds )] , 'v')
    plt.show()

    return linear_spline_func

for i in range(4):
    max_val_min = random.randint(1, 59)
    print(max_val_min)
    spline = get_accumulation_function2(60*60, 60*max_val_min, 10000)
    print(spline(60*20))
    print(spline(60*30))


####################
#
# url = 'https://www.youtube.com/watch?v=4V6pV40RF9A'
# # url = 'https://www.youtube.com/watch?v=rfzlV6oCbG0'
#
# video = pafy.new(url)
# print(f'''
# {video.title}
# {video.duration}
# {video.viewcount}
# ''')
############333
from profanity_filter import ProfanityFilter

pf = ProfanityFilter()

clear_string = 'that is an amazing fight'
bad_string = 'this is some bullshit'

for phrase in [clear_string, bad_string]:
    if pf.is_profane(phrase):
        print(f'{phrase} is PROFANE')

from  pprint import pprint

@dataclass
class Car:
    id: int
    number: str
    mark: str
    owners: typing.List[str]

car = Car(id=1, number='123', mark='Solaris', owners=['Lua', 'Lia'])

pprint(animals)
print(animals)

pprint({'': 1, 'list': ['sdf', 'wef', 234, {"qwre": 15}]})
print({'': 1, 'list': ['sdf', 'wef', 234, {"qwre": 15}]})

pprint(json.dumps({'': 1, 'list': ['sdf', 'wef', 234, {"qwre": 15}]}))
print(json.dumps({'': 1, 'list': ['sdf', 'wef', 234, {"qwre": 15}]}))

pprint(car)
print(car)

pprint(asdict(car))
print(asdict(car))

## не вижу разницы
